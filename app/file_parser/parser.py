"""Parse uploaded files (PDF, DOCX, TXT, …) into plain text.

Usage::

    from app.file_parser.parser import parse_file

    with open("report.pdf", "rb") as f:
        text = parse_file(f.read(), "report.pdf", "application/pdf")
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ── character limit ────────────────────────────────────────────────────
# Keep context-window usage under control.
MAX_FILE_CHARS = 30_000


def parse_file(content: bytes, filename: str, mime: str = "") -> str:
    """Extract readable text from an uploaded file.

    Supported formats:
        .txt   — UTF-8 text (other encodings fall back to replace)
        .pdf   — Portable Document Format (via pypdf)
        .docx  — Word Open XML (via python-docx)
        .md    — Markdown (same as TXT)
        .py / .js / .ts / .json / .csv / .html / .css / .yaml / .xml
               — source code / structured text (same as TXT)

    Args:
        content: Raw file bytes.
        filename: Original file name (used to detect extension).
        mime: Optional MIME type (used as fallback when extension is
              ambiguous).

    Returns:
        Extracted plain text, truncated to ``MAX_FILE_CHARS``.

    Raises:
        ValueError: Unsupported file type (no parser available).
    """
    ext = _get_extension(filename, mime)
    text = _dispatch(ext, content, filename)
    return _truncate(text, filename)


# ── internal helpers ───────────────────────────────────────────────────


def _get_extension(filename: str, mime: str) -> str:
    """Return a file-extension key (lowercase, without dot).

    Tries filename first; falls back to a short mapping of common MIME
    types.
    """
    stem = Path(filename).suffix.lower()
    if stem:
        return stem.lstrip(".")

    MIME_MAP = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "md",
    }
    return MIME_MAP.get(mime, "")


def _dispatch(ext: str, content: bytes, filename: str) -> str:
    """Route to the correct parser based on file extension."""
    if ext in {"txt", "md", "py", "js", "ts", "json", "csv", "html", "css", "yaml", "xml", "ini", "cfg", "conf", "log", "sh", "bat", "ps1", "env", "dockerfile", "gitignore"}:
        return _parse_txt(content)
    if ext == "pdf":
        return _parse_pdf(content)
    if ext == "docx":
        return _parse_docx(content)
    logger.warning("Unsupported file type '.%s' for '%s'", ext, filename)
    return _fallback(content, ext)


def _parse_txt(content: bytes) -> str:
    """Decode bytes as UTF-8 text."""
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF via pypdf."""
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]
    except ImportError:
        logger.error("pypdf not installed — cannot parse PDF")
        return "[PDF 解析失败：缺少 pypdf 库]"

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        logger.warning("PDF parse error: %s", exc)
        return f"[PDF 解析失败：文件格式错误或已损坏]"

    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
        except Exception:
            text = None
        if text:
            pages.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
    return "\n\n".join(pages) if pages else "[PDF 中未提取到文本内容]"


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX via python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        logger.error("python-docx not installed — cannot parse DOCX")
        return "[Word 文档解析失败：缺少 python-docx 库]"

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        logger.warning("DOCX parse error: %s", exc)
        return f"[Word 文档解析失败：文件格式错误或已损坏]"

    sections = [p.text for p in doc.paragraphs if p.text.strip()]
    for index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"--- 表格 {index} ---\n" + "\n".join(rows))
    return "\n\n".join(sections) if sections else "[Word 文档中未提取到文本内容]"


def _fallback(content: bytes, ext: str) -> str:
    """Last-resort: try UTF-8 decode with a warning."""
    text = content.decode("utf-8", errors="replace")
    if not text.strip():
        return f"[不支持的文件格式 .{ext}，且无法作为文本读取]"
    return (
        text
        if len(text) < 200
        else text[:200] + f"\n\n... [文件格式 .{ext} 不是支持的文档类型，仅显示前 200 字符]"
    )


def _truncate(text: str, filename: str) -> str:
    """Cut text to MAX_FILE_CHARS, appending a notice if truncated."""
    if len(text) <= MAX_FILE_CHARS:
        return text
    logger.info("Truncated '%s' at %d chars (was %d)", filename, MAX_FILE_CHARS, len(text))
    return (
        text[:MAX_FILE_CHARS]
        + f"\n\n[... 文件过长，已截断至前 {MAX_FILE_CHARS} 字符。如需查看更多内容，请分段上传]"
    )
