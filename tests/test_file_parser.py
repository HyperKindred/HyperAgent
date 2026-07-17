"""Tests for app/file_parser/parser.py."""

import io

import pytest

from app.file_parser.parser import parse_file, MAX_FILE_CHARS


class TestParseTxt:
    def test_plain_text(self):
        content = "Hello, world!第二行".encode("utf-8")  # 第二行
        result = parse_file(content, "test.txt")
        assert "Hello, world!" in result
        assert "第二行" in result

    def test_markdown(self):
        content = b"# Title\n\nSome **bold** text."
        result = parse_file(content, "readme.md")
        assert "# Title" in result

    def test_code_file(self):
        content = b"def hello():\n    print('hi')"
        result = parse_file(content, "script.py")
        assert "def hello()" in result


class TestParsePdf:
    def test_empty_pdf(self):
        """Create a minimal PDF with pypdf and verify no crash."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)

        result = parse_file(buf.read(), "empty.pdf")
        assert isinstance(result, str)

    def test_non_pdf_bytes(self):
        """Passing non-PDF bytes should not crash."""
        result = parse_file(b"not a pdf at all", "fake.pdf")
        assert isinstance(result, str)


class TestParseDocx:
    def test_non_docx_bytes(self):
        """Passing non-DOCX bytes should not crash."""
        result = parse_file(b"not a docx", "file.docx")
        assert isinstance(result, str)

    def test_extracts_table_cells(self):
        from docx import Document

        document = Document()
        document.add_paragraph("项目进度")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "任务"
        table.cell(0, 1).text = "状态"
        table.cell(1, 0).text = "模型迁移"
        table.cell(1, 1).text = "完成"
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_file(buffer.getvalue(), "进度.docx")

        assert "项目进度" in result
        assert "任务 | 状态" in result
        assert "模型迁移 | 完成" in result


class TestTruncation:
    def test_truncates_long_text(self):
        long = b"A" * (MAX_FILE_CHARS + 1000)
        result = parse_file(long, "long.txt")
        assert len(result) <= MAX_FILE_CHARS + 200  # truncation notice
        assert "截断" in result  # 截断

    def test_short_text_not_truncated(self):
        short = b"Hello"
        result = parse_file(short, "short.txt")
        assert result == "Hello"


class TestUnsupported:
    def test_unknown_extension_fallback(self):
        result = parse_file(b"some content", "archive.zip")
        assert isinstance(result, str)

    def test_empty_file(self):
        result = parse_file(b"", "empty.txt")
        assert result == ""


class TestEdgeCases:
    def test_no_extension(self):
        result = parse_file(b"hello", "README")
        assert "hello" in result

    def test_invalid_utf8(self):
        raw = b"\x80\x81\x82"
        result = parse_file(raw, "binary.bin")
        assert isinstance(result, str)

    def test_mime_fallback(self):
        result = parse_file(b"plain text", "unknown", mime="text/plain")
        assert "plain text" in result

    def test_mime_unrecognised(self):
        result = parse_file(b"data", "unknown", mime="application/octet-stream")
        assert isinstance(result, str)
